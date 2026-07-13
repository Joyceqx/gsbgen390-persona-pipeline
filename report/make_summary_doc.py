"""Generate the GSBGEN390 project summary as a Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/Users/joyce/Developer/gsbgen390/notes/GSBGEN390_Project_Summary.docx")
OUT.parent.mkdir(exist_ok=True)
INK = RGBColor(0x22, 0x22, 0x22)
ACCENT = RGBColor(0x2f, 0x6f, 0x9f)
MUTED = RGBColor(0x5f, 0x6b, 0x76)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = INK
for m in (doc.sections[0].top_margin, doc.sections[0].bottom_margin):
    pass
doc.sections[0].top_margin = Inches(0.8)
doc.sections[0].bottom_margin = Inches(0.8)
doc.sections[0].left_margin = Inches(0.9)
doc.sections[0].right_margin = Inches(0.9)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = INK
    p.space_after = Pt(2)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)


def h(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)


def para(text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    # support **bold** segments
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2]); r.font.bold = True
        else:
            p.add_run(seg)
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if lead:
        r = p.add_run(lead + " "); r.font.bold = True
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2]); r.font.bold = True
        else:
            p.add_run(seg)


# ---------------------------------------------------------------------------
title("GSBGEN390 — Project Summary (Phase 1)")
subtitle("Joyce Yu · Stanford GSB · Advisor: Prof. Mohsen Bayati · 2 June 2026")

h("1. What the project is")
para("This is a methodological study of how well large language models (LLMs) can "
     "stand in for survey respondents. We give an LLM a person's answers to the 2024 "
     "General Social Survey (GSS) as a “persona,” then ask it to predict that "
     "person's answers to a set of held-out attitude questions. The central question is "
     "**which categories of background information — demographic, behavioral, "
     "psychological, or attitudinal — actually drive the model's prediction.** The "
     "work is benchmarked against Park et al. (2024), the leading study of LLM persona "
     "simulation.")

h("2. Research design")
bullet("the 2024 GSS cross-section (about 3,300 respondents). Public data, no privacy constraints.", lead="Data:")
bullet("predict 12 held-out attitude questions (political ideology and party, abortion, "
       "death penalty, guns, two gender-role items, racial attitudes, confidence in banks "
       "and Congress, redistribution, financial satisfaction) from each respondent's other GSS answers.", lead="Task:")
bullet("the persona is built from four feature categories — demographic, behavioral, "
       "psychological, attitudinal. To prevent leakage, when predicting a question we drop its "
       "whole topic battery from the persona.", lead="Features & leakage control:")
bullet("five candidates — four cheap LLMs (Qwen3-Max, DeepSeek-V3.1, Llama-4-Maverick, "
       "Kimi-K2) plus a Random policy that assigns one of the four per respondent — and "
       "GPT-4o as an expensive reference (“anchor”).", lead="Models:")
bullet("three published persona-prompt formats: P0 key-value list (Park), P1 first-person "
       "prose (Argyle), P2 interview Q&A (Wang).", lead="Prompts:")
bullet("normalized error (distance scaled by the question's range, 0–1) as the primary "
       "score, with exact-match accuracy as a secondary number; compared against a base-rate "
       "guess and a no-persona baseline.", lead="Metrics:")
bullet("Phase 1A selects the model and prompt (done). Phase 1B runs the chosen setup at full "
       "scale and measures each feature category's contribution by removing it and seeing how "
       "much the prediction moves. Phase 1C decomposes this to the battery level. Phase 2 (later) "
       "extends to personality and behavioral-economic outcomes.", lead="Phases:")

h("3. Literature review")
para("The review spans six themes, grounding both the method and the evaluation:")
bullet("**Foundational work** — Park et al. (2023, 2024) on generative agents; the 2024 "
       "paper is our benchmark and the source of the survey-vs-interview accuracy gap we are testing.")
bullet("**Persona-prompt formats** — Argyle et al. (2023), “Out of One, Many”, and "
       "Wang et al. (2025), “The Prompt Makes the Persona”; these justify our three prompt formats.")
bullet("**Evaluation & effect sizes** — PersonaGym, Eval4Sim, and Funder & Ozer (2019) for "
       "judging when an effect is substantively meaningful rather than just significant.")
bullet("**Validity skepticism** — Bisbee et al. and others on the perils of treating LLM "
       "output as survey data, and on social-desirability bias; this is the counter-evidence we hold ourselves against.")
bullet("**Judgment-accuracy theory** — Brunswik's Lens Model, Funder's Realistic Accuracy "
       "Model, and Vazire's self–other knowledge asymmetry, which frame what “accurate "
       "prediction from cues” means.")
bullet("**Construct theory** for the question batteries — Converse on belief systems, Moral "
       "Foundations, trust and well-being scales — so the eval items are theoretically coherent, not ad hoc.")

h("4. Current progress and what we have learned")
para("Phase 1A is complete. We ran the full panel (four models × three prompts × 12 "
     "questions × 200 respondents), the GPT-4o anchor (100 respondents, two leakage settings), "
     "and a no-persona baseline. Prof. Bayati independently reanalyzed the predictions; the findings below incorporate his review.")
bullet("**Normalized error is the right metric, not exact match.** Being one step off on a "
       "7-point scale is not the same as missing a yes/no. The per-question ranking actually "
       "flips between the two scores, so we have made normalized error primary.")
bullet("**No single model beats a random mixture of them.** With standard errors clustered by "
       "respondent, no model is significantly better than the Random policy on either metric; "
       "only Qwen separates, and it is worse. The model choice does not move the score.")
bullet("**Mode collapse is the one real difference.** On the two institutional-trust questions "
       "(confidence in banks, confidence in Congress), every model except Kimi gives essentially "
       "the same answer to all 200 respondents — it ignores the persona. The metrics reward "
       "this, but it matters for Phase 1B, where we read feature importance from how much the prediction moves.")
bullet("**The cheap model is good enough.** Under the same rules, GPT-4o is a statistical tie "
       "with the cheap models (a coin flip, p = 0.41). Paying roughly seven times more per call "
       "buys no measurable accuracy. This is the result that licenses scaling on a cheap model.")
bullet("**The persona adds real but modest signal.** Against a fair no-persona baseline (the same "
       "model with no persona) it helps on 9 of 12 questions; the clear wins are party ID and "
       "abortion, while on several questions it only matches, or slightly trails, a base-rate guess.")
bullet("**Prompt is the one clear lever.** P1 and P2 both significantly beat P0 on both metrics; "
       "P1 and P2 are about even with each other.")
bullet("**A per-question router (Prof. Bayati's idea)** gives a small, real gain on normalized "
       "error, but its learned policy sends the collapse-prone questions to the collapsing models. "
       "We keep it as a secondary result.")

h("5. Open questions and next steps")
bullet("**Model choice is an open decision** for the advisor: by the metrics any model (including "
       "Random) is fine, but the mode-collapse concern argues for Kimi, the only model that keeps "
       "every question responsive to the persona. We have laid out both sides rather than concluding.")
bullet("**Prompt:** lean P2 (marginally best on normalized error; P1 equally defensible).")
bullet("**Phase 1B:** run the chosen setup at full scale with leave-one-out feature ablations to "
       "produce the attribution, and include a no-persona baseline so the attribution has a fair zero point.")
bullet("**Park comparison** stays at the aggregate level; Park published no per-question "
       "survey-only table, so a per-question side-by-side would be misleading.")

h("Deliverables to date")
para("Full Phase 1A report with figures (HTML); the analysis databook (Excel) and raw prediction "
     "data (CSV); the reanalysis scripts; and a drafted reply to Prof. Bayati. All reproducible "
     "from the raw predictions.", after=2)

doc.save(str(OUT))
print("wrote", OUT)
print("pages (approx):", len(doc.paragraphs), "paragraphs")
